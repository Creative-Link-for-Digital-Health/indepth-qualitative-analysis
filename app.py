"""
Qualitative Coding Analysis Application
A Streamlit app for analyzing interview transcripts using a hierarchical codebook.
"""

import streamlit as st
import toml
import os
import re
import json
import docx
import PyPDF2
from typing import Dict, List, Tuple, Optional
from openai import OpenAI
from io import BytesIO
from pathlib import Path

# ==================== API UTILITIES ====================

def load_api_params(secrets_path: str) -> Dict[str, str]:
    """Load API parameters from a TOML file."""
    try:
        with open(secrets_path, 'r') as f:
            secrets = toml.load(f)
        return {
            'API_KEY': secrets.get('API_KEY', 'ollama'),
            'API_URL': secrets.get('API_URL', 'http://localhost:11434/v1'),
            'MODEL': secrets.get('MODEL', 'llama3.1'),
        }
    except FileNotFoundError:
        st.error(f"Error: {secrets_path} file not found")
        st.info(f"Please create a {secrets_path} file with your API settings.")
        return {
            'API_KEY': 'ollama',
            'API_URL': 'http://localhost:11434/v1',
            'MODEL': 'llama3.1'
        }
    except Exception as e:
        st.error(f"Error loading API parameters: {e}")
        return {
            'API_KEY': 'ollama',
            'API_URL': 'http://localhost:11434/v1',
            'MODEL': 'llama3.1'
        }

def initialize_client(api_params: Dict[str, str]) -> OpenAI:
    """Initialize the OpenAI client with the given configuration."""
    try:
        return OpenAI(
            api_key=api_params['API_KEY'],
            base_url=api_params['API_URL']
        )
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {e}")
        return None

def generate_completion(client: OpenAI, model: str, messages: List[Dict], temperature: float = 0.3) -> str:
    """Generate completion using the OpenAI API."""
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating completion: {e}")
        return f"Error: {str(e)}"

# ==================== FILE PARSING UTILITIES ====================

def parse_file(file) -> Optional[str]:
    """Parse file content based on file extension."""
    file_extension = os.path.splitext(file.name)[1].lower()
    
    try:
        if file_extension == '.txt':
            return file.getvalue().decode('utf-8')
        elif file_extension == '.docx':
            doc = docx.Document(file)
            return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
        elif file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        print(f"Error parsing file {file.name}: {e}")
        return None

def prepare_transcript_with_ids(transcript_text: str) -> Tuple[str, Dict[str, str]]:
    """Process the transcript to add IDs to each paragraph for hyperlinking."""
    lines = transcript_text.split('\n')
    processed_lines = []
    quote_map = {}
    
    for i, line in enumerate(lines):
        line_id = f"line-{i}"
        processed_lines.append(f'<p id="{line_id}">{line}</p>')
        
        # Store the line with its ID for later quote matching
        if line.strip():
            quote_map[line.strip()] = line_id
    
    return '\n'.join(processed_lines), quote_map

def parse_codebook(codebook_text: str) -> dict:
    """Parse the codebook text into a structured format."""
    lines = codebook_text.split('\n')
    codebook = {
        'title': '',
        'parent_codes': []
    }
    
    current_parent = None
    current_child = None
    current_grandchild = None
    current_great_grandchild = None
    
    current_section = None
    current_content = []
    
    # Extract title from the first line if it starts with #
    if lines and lines[0].startswith('# '):
        codebook['title'] = lines[0][2:].strip()
    
    for line in lines:
        # Skip empty lines
        if not line.strip():
            continue
            
        # Parent code (e.g., ## [Parent Code 1] Coping Strategies)
        if line.startswith('## [Parent Code'):
            # Save previous section if exists
            if current_section and current_content:
                _save_current_section(current_parent, current_child, current_grandchild, 
                                     current_great_grandchild, current_section, current_content)
            
            # Extract code number and name
            parts = line[3:].split(']', 1)
            code_id = parts[0].strip()
            name = parts[1].strip() if len(parts) > 1 else ""
            
            current_parent = {
                'id': code_id,
                'name': name,
                'children': []
            }
            codebook['parent_codes'].append(current_parent)
            current_child = None
            current_grandchild = None
            current_great_grandchild = None
            current_section = None
            current_content = []
            
        # Child code (e.g., ### [Child Code 1.1] Approach Coping)
        elif line.startswith('### [Child Code'):
            # Save previous section if exists
            if current_section and current_content:
                _save_current_section(current_parent, current_child, current_grandchild, 
                                     current_great_grandchild, current_section, current_content)
            
            # Extract code number and name
            parts = line[4:].split(']', 1)
            code_id = parts[0].strip()
            name = parts[1].strip() if len(parts) > 1 else ""
            
            current_child = {
                'id': code_id,
                'name': name,
                'children': []
            }
            if current_parent:
                current_parent['children'].append(current_child)
            current_grandchild = None
            current_great_grandchild = None
            current_section = None
            current_content = []
            
        # Grandchild code (e.g., #### [Grandchild Code 1.1.1] Adaptive Coping)
        elif line.startswith('#### [Grandchild Code'):
            # Save previous section if exists
            if current_section and current_content:
                _save_current_section(current_parent, current_child, current_grandchild, 
                                     current_great_grandchild, current_section, current_content)
            
            # Extract code number and name
            parts = line[5:].split(']', 1)
            code_id = parts[0].strip()
            name = parts[1].strip() if len(parts) > 1 else ""
            
            current_grandchild = {
                'id': code_id,
                'name': name,
                'children': []
            }
            if current_child:
                current_child['children'].append(current_grandchild)
            current_great_grandchild = None
            current_section = None
            current_content = []
            
        # Great grandchild code (e.g., ##### [Great Grandchild Code 1.1.1.1] Caregiver)
        elif line.startswith('##### [Great Grandchild Code'):
            # Save previous section if exists
            if current_section and current_content:
                _save_current_section(current_parent, current_child, current_grandchild, 
                                     current_great_grandchild, current_section, current_content)
            
            # Extract code number and name
            parts = line[6:].split(']', 1)
            code_id = parts[0].strip()
            name = parts[1].strip() if len(parts) > 1 else ""
            
            current_great_grandchild = {
                'id': code_id,
                'name': name
            }
            if current_grandchild:
                current_grandchild['children'].append(current_great_grandchild)
            current_section = None
            current_content = []
            
        # Section headers (e.g., **Definition:** or **Inclusion criteria:**)
        elif line.strip().startswith('**') and line.strip().endswith(':**'):
            # Save previous section if exists
            if current_section and current_content:
                _save_current_section(current_parent, current_child, current_grandchild, 
                                     current_great_grandchild, current_section, current_content)
            
            # Extract section name (remove ** and :** from section name)
            current_section = line.strip()[2:-3]  
            current_content = []
            
        # Section content
        elif current_section is not None:
            current_content.append(line.strip())
    
    # Save the last section if exists
    if current_section and current_content:
        _save_current_section(current_parent, current_child, current_grandchild, 
                             current_great_grandchild, current_section, current_content)
    
    return codebook

def _save_current_section(parent, child, grandchild, great_grandchild, section_name, content):
    """Helper function to save the current section content to the appropriate code level."""
    section_content = '\n'.join(content)
    
    if great_grandchild is not None:
        great_grandchild[section_name] = section_content
    elif grandchild is not None:
        grandchild[section_name] = section_content
    elif child is not None:
        child[section_name] = section_content
    elif parent is not None:
        parent[section_name] = section_content

def markdown_to_docx(markdown_text: str) -> bytes:
    """Convert markdown text to a DOCX file."""
    # Create a new document
    doc = docx.Document()
    
    # Process the markdown text line by line
    lines = markdown_text.split('\n')
    i = 0
    
    while i < len(lines):
        line = lines[i].rstrip()
        
        # Skip empty lines
        if not line:
            i += 1
            continue
        
        # Handle headings
        if line.startswith('# '):
            # Heading 1
            doc.add_heading(line[2:], level=1)
        elif line.startswith('## '):
            # Heading 2
            doc.add_heading(line[3:], level=2)
        elif line.startswith('### '):
            # Heading 3
            doc.add_heading(line[4:], level=3)
        elif line.startswith('#### '):
            # Heading 4
            doc.add_heading(line[5:], level=4)
        elif line.startswith('##### '):
            # Heading 5
            doc.add_heading(line[6:], level=5)
        elif line.startswith('- '):
            # Bullet point
            p = doc.add_paragraph(line[2:], style='List Bullet')
            
            # Check if this is a quote with explanation
            if '**Quote**:' in line:
                # Format the quote line
                p.text = p.text.replace('**Quote**:', 'Quote:')
                p.runs[0].bold = True
                
                # Check for explanation in the next lines
                if i + 2 < len(lines) and '**Explanation**:' in lines[i + 2]:
                    # Add a blank line
                    doc.add_paragraph()
                    
                    # Add explanation
                    exp_line = lines[i + 2]
                    exp_p = doc.add_paragraph(exp_line.replace('**Explanation**:', 'Explanation:'))
                    exp_p.runs[0].bold = True
                    
                    # Skip the processed lines
                    i += 2
        else:
            # Regular paragraph
            p = doc.add_paragraph()
            
            # Process bold text (**text**)
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    # Bold text
                    run = p.add_run(part[2:-2])
                    run.bold = True
                else:
                    # Regular text
                    p.add_run(part)
        
        i += 1
    
    # Save the document to a BytesIO object
    docx_bytes = BytesIO()
    doc.save(docx_bytes)
    docx_bytes.seek(0)
    
    return docx_bytes.getvalue()

# ==================== ANALYSIS UTILITIES ====================

def analyze_transcript_for_themes(client, model: str, transcript_text: str, 
                                codebook: Dict, quote_map: Dict) -> str:
    """Use AI to analyze the transcript for themes and extract relevant quotes."""
    system_prompt = {
        "role": "system", 
        "content": f"""You are an expert qualitative researcher analyzing interview transcripts. 
        
        Your task is to extract quotes from the transcript that match the codes from the codebook.
        
        Here is the codebook:
        {json.dumps(codebook, indent=2)}
        
        The codebook has a hierarchical structure:
        - Parent Codes: Main categories (e.g., [Parent Code 1])
        - Child Codes: Subcategories under parent codes (e.g., [Child Code 1.1])
        - Grandchild Codes: Further subcategories (e.g., [Grandchild Code 1.1.1])
        - Great Grandchild Codes: The most specific level (e.g., [Great Grandchild Code 1.1.1.1])
        
        For each code at any level:
        1. Find ALL relevant quotes from the transcript that match this code
        2. For each quote, provide the exact text and a brief explanation of why it matches
        
        Organize your response hierarchically by parent code, then child code, etc. Use bullet points for clarity.
        Make sure each quote is directly relevant to the code it's matched with.
        
        Format your response as:
        
        # [Parent Code 1]: [Parent Code Name]
        
        ## [Child Code 1.1]: [Child Code Name]
        - **Quote**: "[Exact quote from transcript]"
          
          **Explanation**: Brief explanation of why this quote matches the code
        
        - **Quote**: "[Another relevant quote]"
          
          **Explanation**: Brief explanation
        
        ### [Grandchild Code 1.1.1]: [Grandchild Code Name]
        - **Quote**: "[Exact quote from transcript]"
          
          **Explanation**: Brief explanation
        
        # [Parent Code 2]: [Parent Code Name]
        
        And so on for all codes.
        
        IMPORTANT: 
        1. Make sure to put the explanation on a separate line from the quote, with a blank line between them. 
        2. Both "Quote" and "Explanation" should be in bold.
        3. Include ALL relevant quotes from the transcript, even if there are many for a single code.
        4. Use exact quotes from the transcript - do not paraphrase or modify them in any way.
        5. Quotes must match text from the transcript exactly for hyperlinking to work correctly.
        6. Maintain the hierarchical structure in your response to match the codebook structure.
        """
    }
    
    user_message = {
        "role": "user",
        "content": f"Please analyze the following interview transcript and extract relevant quotes for each theme and code:\n\n{transcript_text}"
    }
    
    return generate_completion(
        client=client,
        model=model,
        messages=[system_prompt, user_message]
    )

def add_permalink_icons(analysis_text: str, quote_map: Dict[str, str]) -> str:
    """Add permalink icon anchors to quotes in the analysis results."""
    lines = analysis_text.split('\n')
    result_lines = []
    quote_counter = 0
    
    i = 0
    while i < len(lines):
        line = lines[i]
        # Handle parent code headers (H1)
        if line.startswith('# '):
            parent_id = line[2:].strip().lower().replace(' ', '-').replace(':', '')
            result_lines.append(f'{line} <a href="#{parent_id}" class="header-anchor">🔗</a>')
        
        # Handle child code headers (H2)
        elif line.startswith('## '):
            child_id = line[3:].strip().lower().replace(' ', '-').replace(':', '')
            result_lines.append(f'{line} <a href="#{child_id}" class="header-anchor">🔗</a>')
            
        # Handle grandchild code headers (H3)
        elif line.startswith('### '):
            grandchild_id = line[4:].strip().lower().replace(' ', '-').replace(':', '')
            result_lines.append(f'{line} <a href="#{grandchild_id}" class="header-anchor">🔗</a>')
            
        # Handle great grandchild code headers (H4)
        elif line.startswith('#### '):
            great_grandchild_id = line[5:].strip().lower().replace(' ', '-').replace(':', '')
            result_lines.append(f'{line} <a href="#{great_grandchild_id}" class="header-anchor">🔗</a>')
        
        # Handle quote lines
        elif line.strip().startswith('- '):
            quote_line = line
            
            if "**Quote**:" in quote_line:
                quote_match = re.search(r'\*\*Quote\*\*: "(.*?)"', quote_line)
                if quote_match:
                    quote_text = quote_match.group(1)
                    quote_counter += 1
                    quote_id = f'quote-{quote_counter}'
                    
                    # Find matching transcript line ID
                    target_id = None
                    for transcript_line, line_id in quote_map.items():
                        if quote_text in transcript_line:
                            target_id = line_id
                            break
                    
                    # Format quote with proper styling
                    modified_line = quote_line.replace(
                        '**Quote**:', 
                        '<span class="quote-text-bold">Quote</span>:'
                    )
                    
                    # Add permalink container
                    if target_id:
                        modified_line = (
                            f'<div id="{quote_id}" class="quote-container">'
                            f'{modified_line} '
                            f'<a href="#{target_id}" class="quote-anchor">🔗</a>'
                            f'</div>'
                        )
                    else:
                        modified_line = (
                            f'<div id="{quote_id}" class="quote-container">'
                            f'{modified_line}'
                            f'</div>'
                        )
                    
                    result_lines.append(modified_line)
                else:
                    result_lines.append(quote_line)
                
                # Handle explanation if present
                if i + 2 < len(lines) and "**Explanation**:" in lines[i + 2]:
                    explanation_line = lines[i + 2].replace(
                        '**Explanation**:',
                        '<span class="quote-text-bold">Explanation</span>:'
                    )
                    result_lines.append(lines[i + 1])  # Add blank line
                    result_lines.append(explanation_line)
                    i += 2  # Skip processed lines
            else:
                result_lines.append(line)
        else:
            result_lines.append(line)
        
        i += 1
    
    return '\n'.join(result_lines)

# ==================== UI COMPONENTS ====================

def render_codebook(codebook: Dict):
    """Render the codebook section with hierarchical structure."""
    st.header("Codebook")
    if 'title' in codebook and codebook['title']:
        st.subheader(codebook['title'])
    
    # Render parent codes
    for parent in codebook['parent_codes']:
        with st.expander(f"{parent['id']} {parent['name']}"):
            # Display parent code sections if they exist
            _render_code_sections(parent)
            
            # Render child codes using containers instead of nested expanders
            if 'children' in parent and parent['children']:
                for child in parent['children']:
                    st.markdown(f"### ↳ {child['id']} {child['name']}")
                    child_container = st.container()
                    
                    # Display child code sections with indentation
                    with child_container:
                        st.markdown('<div style="margin-left: 20px;">', unsafe_allow_html=True)
                        _render_code_sections(child)
                        
                        # Render grandchild codes with further indentation
                        if 'children' in child and child['children']:
                            for grandchild in child['children']:
                                st.markdown(f"#### ↳↳ {grandchild['id']} {grandchild['name']}")
                                grandchild_container = st.container()
                                
                                # Display grandchild code sections with more indentation
                                with grandchild_container:
                                    st.markdown('<div style="margin-left: 20px;">', unsafe_allow_html=True)
                                    _render_code_sections(grandchild)
                                    
                                    # Render great grandchild codes with even more indentation
                                    if 'children' in grandchild and grandchild['children']:
                                        for great_grandchild in grandchild['children']:
                                            st.markdown(f"##### ↳↳↳ {great_grandchild['id']} {great_grandchild['name']}")
                                            great_grandchild_container = st.container()
                                            
                                            # Display great grandchild code sections with the most indentation
                                            with great_grandchild_container:
                                                st.markdown('<div style="margin-left: 20px;">', unsafe_allow_html=True)
                                                _render_code_sections(great_grandchild)
                                                st.markdown('</div>', unsafe_allow_html=True)
                                    
                                    st.markdown('</div>', unsafe_allow_html=True)
                        
                        st.markdown('</div>', unsafe_allow_html=True)

def _render_code_sections(code: Dict):
    """Helper function to render code sections like Definition, Inclusion criteria, etc."""
    # List of common section names to check for
    sections = [
        "Definition", "Inclusion criteria", "Exclusion criteria", 
        "Example", "Examples", "Notes", "Description"
    ]
    
    # Display any sections that exist in the code
    for section in sections:
        if section in code:
            st.markdown(f"**{section}:** {code[section]}")
    
    # Display any other sections not in our predefined list
    for key, value in code.items():
        if key not in ['id', 'name', 'children'] and key not in sections:
            st.markdown(f"**{key}:** {value}")

# ==================== CUSTOM STYLING ====================

CUSTOM_CSS = """
/* Transcript container that respects theme */
.transcript-container {
    height: 600px;
    overflow-y: auto;
    border: 1px solid var(--text-color);
    padding: 15px;
    font-family: monospace;
    white-space: pre-wrap;
    background-color: var(--background-color);
    border-radius: 5px;
    color: var(--text-color);
}

/* Dark mode variables */
[data-theme="dark"] {
    --background-color: #262730;
    --text-color: #fafafa;
    --highlight-color: #555555;
}

/* Light mode variables */
[data-theme="light"] {
    --background-color: #ffffff;
    --text-color: #31333F;
    --highlight-color: #f0f0f0;
}

.transcript-container p {
    margin: 0;
    padding: 2px 0;
}

.transcript-container p:target {
    background-color: #FFFF00;
    padding: 2px 5px;
    border-radius: 3px;
    color: #000000;
}

/* Style for permalink icon */
.header-anchor {
    opacity: 0;
    transition: opacity 0.2s ease-in-out;
    text-decoration: none;
    margin-left: 8px;
    font-size: 0.8em;
}

h1:hover .header-anchor,
h2:hover .header-anchor,
h3:hover .header-anchor,
h4:hover .header-anchor,
.quote-container:hover .quote-anchor {
    opacity: 1;
}

.quote-container {
    position: relative;
    padding-right: 25px;
}

.quote-anchor {
    opacity: 0;
    position: absolute;
    right: 0;
    top: 0;
    text-decoration: none;
}

/* Fix for properly displaying bold text */
.quote-text-bold {
    font-weight: bold;
}
"""

THEME_SCRIPT = """
<script>
// Script to detect and set theme
document.addEventListener('DOMContentLoaded', function() {
    // Check if the page has a dark background
    const isDarkMode = window.matchMedia && window.matchMedia('(prefers-color-scheme: dark)').matches;
    document.body.setAttribute('data-theme', isDarkMode ? 'dark' : 'light');
    
    // Listen for theme changes
    window.matchMedia('(prefers-color-scheme: dark)').addEventListener('change', e => {
        document.body.setAttribute('data-theme', e.matches ? 'dark' : 'light');
    });
});
</script>
"""

# ==================== MAIN APPLICATION ====================

def initialize_session_state():
    """Initialize or reset session state variables."""
    if 'codebook' not in st.session_state:
        st.session_state.codebook = None
    if 'transcript_text' not in st.session_state:
        st.session_state.transcript_text = None
    if 'processed_transcript' not in st.session_state:
        st.session_state.processed_transcript = None
    if 'quote_map' not in st.session_state:
        st.session_state.quote_map = {}
    if 'analysis_results' not in st.session_state:
        st.session_state.analysis_results = None
    if 'transcript_filename' not in st.session_state:
        st.session_state.transcript_filename = "transcript"
    
    # Load API parameters
    if 'api_params' not in st.session_state:
        st.session_state.api_params = load_api_params('.secrets.toml')
    
    # Initialize OpenAI client
    if 'client' not in st.session_state:
        st.session_state.client = initialize_client(st.session_state.api_params)

def main():
    """Main application entry point."""
    # Configure page
    st.set_page_config(
        page_title="In-Depth Qual Coding",
        page_icon="📊",
        layout="wide"
    )
    
    # Initialize session state
    initialize_session_state()
    
    # Apply custom styling
    st.markdown(f"<style>{CUSTOM_CSS}</style>", unsafe_allow_html=True)
    st.markdown(THEME_SCRIPT, unsafe_allow_html=True)
    
    # Display title
    st.title("Qualitative Coding Analysis")
    st.write("Upload your codebook and interview transcripts to extract themes and relevant quotes.")
    
    # Sidebar
    with st.sidebar:
        st.header("Upload Files")
        
        # Codebook upload
        codebook_file = st.file_uploader(
            "Upload Codebook (txt, docx, pdf)", 
            type=["txt", "docx", "pdf"], 
            key="codebook_uploader"
        )
        
        if codebook_file is not None and (
            'codebook' not in st.session_state or 
            st.button("Process Codebook")
        ):
            with st.spinner("Processing codebook..."):
                codebook_text = parse_file(codebook_file)
                if codebook_text:
                    st.session_state.codebook = parse_codebook(codebook_text)
                    st.success("Codebook processed successfully!")
        
        # Transcript upload
        transcript_file = st.file_uploader(
            "Upload Transcript (txt, docx, pdf)", 
            type=["txt", "docx", "pdf"], 
            key="transcript_uploader"
        )
        
        if transcript_file is not None and (
            'transcript_text' not in st.session_state or 
            st.button("Process Transcript")
        ):
            with st.spinner("Processing transcript..."):
                transcript_text = parse_file(transcript_file)
                if transcript_text:
                    # Store the original filename (without extension) for later use
                    filename = os.path.splitext(transcript_file.name)[0]
                    st.session_state.transcript_filename = filename
                    
                    st.session_state.transcript_text = transcript_text
                    # Process transcript to add IDs for hyperlinking
                    processed_transcript, quote_map = prepare_transcript_with_ids(transcript_text)
                    st.session_state.processed_transcript = processed_transcript
                    st.session_state.quote_map = quote_map
                    st.success("Transcript processed successfully!")
        
        # Analysis button
        if (st.session_state.get('codebook') and 
            st.session_state.get('transcript_text') and 
            st.button("Analyze Transcript")):
            with st.spinner("Analyzing transcript... This may take a minute."):
                results = analyze_transcript_for_themes(
                    client=st.session_state.client,
                    model=st.session_state.api_params['MODEL'],
                    transcript_text=st.session_state.transcript_text,
                    codebook=st.session_state.codebook,
                    quote_map=st.session_state.quote_map
                )
                if results:
                    st.session_state.analysis_results = results
                    st.success("Analysis complete!")
    
    # Main content area
    # Display codebook if available
    if st.session_state.get('codebook'):
        render_codebook(st.session_state.codebook)
    
    # Display analysis results if available
    if st.session_state.get('analysis_results'):
        col1, col2 = st.columns([3, 2])
        
        with col1:
            st.header("Analysis Results")
            
            # Process and display results with permalink icons
            processed_results = add_permalink_icons(
                st.session_state.analysis_results,
                st.session_state.quote_map
            )
            st.markdown(processed_results, unsafe_allow_html=True)
            
            # Download button for results
            # Convert markdown to DOCX
            docx_data = markdown_to_docx(st.session_state.analysis_results)
            
            # Use the original transcript filename for the download
            filename = f"{st.session_state.transcript_filename}_results.docx"
            
            if st.download_button(
                label="Download Analysis Results",
                data=docx_data,
                file_name=filename,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ):
                st.success(f"Downloaded analysis results as {filename}!")
        
        with col2:
            st.header("Transcript")
            if st.session_state.get('processed_transcript'):
                st.markdown(
                    f'<div class="transcript-container">{st.session_state.processed_transcript}</div>',
                    unsafe_allow_html=True
                )
            else:
                st.write("No transcript processed yet.")

if __name__ == "__main__":
    main()