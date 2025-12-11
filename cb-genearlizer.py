"""
Codebook Generalizer
A Streamlit app for converting various codebook formats into a standardized hierarchical structure.
"""

import streamlit as st
import toml
import os
import re
import json
import docx
import PyPDF2
from typing import Dict, List, Optional
from openai import OpenAI
from io import BytesIO
import pandas as pd

# ==================== API UTILITIES ====================

def load_api_params(secrets_path: str = '.secrets.toml') -> Dict[str, str]:
    """Load API parameters from a TOML file."""
    try:
        with open(secrets_path, 'r') as f:
            secrets = toml.load(f)
        return {
            'API_KEY': secrets.get('API_KEY', 'ollama'),
            'API_URL': secrets.get('API_URL', 'http://localhost:11434/v1'),
            'MODEL': secrets.get('MODEL', 'llama3.1'),
        }
    except FileNotFoundError:
        st.error(f"Error: {secrets_path} file not found")
        st.info(f"Please create a {secrets_path} file with your API settings.")
        return {
            'API_KEY': 'ollama',
            'API_URL': 'http://localhost:11434/v1',
            'MODEL': 'llama3.1'
        }

def initialize_client(api_params: Dict[str, str]) -> OpenAI:
    """Initialize the OpenAI client with the given configuration."""
    try:
        return OpenAI(
            api_key=api_params['API_KEY'],
            base_url=api_params['API_URL']
        )
    except Exception as e:
        st.error(f"Error initializing OpenAI client: {e}")
        return None

def generate_completion(client: OpenAI, model: str, messages: List[Dict], temperature: float = 0.1) -> str:
    """Generate completion using the OpenAI API."""
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            stream=False
        )
        return response.choices[0].message.content
    except Exception as e:
        st.error(f"Error generating completion: {e}")
        return None

# ==================== FILE PARSING UTILITIES ====================

def parse_file(file) -> Optional[str]:
    """Parse file content based on file extension."""
    file_extension = os.path.splitext(file.name)[1].lower()
    
    try:
        if file_extension == '.txt':
            return file.getvalue().decode('utf-8')
        elif file_extension == '.docx':
            doc = docx.Document(file)
            # Extract all text including tables
            full_text = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                full_text.append(paragraph.text)
            
            # Extract tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        row_text.append(cell.text.strip())
                    full_text.append(' | '.join(row_text))
            
            return '\n'.join(full_text)
        elif file_extension == '.pdf':
            pdf_reader = PyPDF2.PdfReader(file)
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() or ""
            return text
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    except Exception as e:
        st.error(f"Error parsing file {file.name}: {e}")
        return None

def clean_text(text: str) -> str:
    """Clean text by removing Word artifacts and excessive whitespace."""
    # Remove Word track changes artifacts
    text = re.sub(r'\{\.comment-start.*?\}', '', text)
    text = re.sub(r'\{\.comment-end.*?\}', '', text)
    text = re.sub(r'\[.*?\]\{\.mark\}', lambda m: m.group(0).replace('{.mark}', ''), text)
    
    # Remove excessive whitespace
    text = re.sub(r'\n\s*\n', '\n\n', text)
    
    return text.strip()

# ==================== CODEBOOK PARSING WITH AI ====================

def detect_and_parse_codebook(client: OpenAI, model: str, raw_text: str) -> Dict:
    """Use AI to detect codebook structure and parse it into standardized format."""
    
    system_prompt = {
        "role": "system",
        "content": """You are an expert at parsing qualitative research codebooks. Your task is to:

1. Identify the hierarchical structure of the codebook
2. Extract all codes with their relationships
3. Standardize the output format

Common patterns to look for:
- "Theme" or "Domain" = Parent Code level
- "Code" under a theme = Child Code level
- Sub-codes (e.g., 1.1.1) = Grandchild Code level
- Further sub-codes = Great Grandchild Code level

For each code, extract:
- The code name/title
- Definition (if present)
- Inclusion criteria (if present)
- Exclusion criteria (if present)
- Examples or exemplar quotes (if present)

Output ONLY valid JSON (no comments, no ellipsis, no markdown) with this exact structure:
{
  "title": "Codebook Title (if identifiable)",
  "parent_codes": [
    {
      "id": "Parent Code 1",
      "name": "Name of the parent code",
      "definition": "Definition text if present",
      "inclusion_criteria": "Inclusion criteria if present",
      "exclusion_criteria": "Exclusion criteria if present",
      "examples": "Examples if present",
      "children": [
        {
          "id": "Child Code 1.1",
          "name": "Name of the child code",
          "definition": "Definition text if present",
          "inclusion_criteria": "Inclusion criteria if present",
          "exclusion_criteria": "Exclusion criteria if present",
          "examples": "Examples if present",
          "children": []
        }
      ]
    }
  ]
}

Important rules:
1. Use standardized IDs: "Parent Code 1", "Child Code 1.1", etc. (NOT "Theme-1")
2. Include ALL codes found in the document
3. Empty fields should be empty strings ""
4. Children arrays should be empty [] if no children exist
5. Do NOT include comments or ellipsis (...) in the JSON
6. Make intelligent guesses about structure when unclear
7. Don't include visual elements like "coding trees"
"""
    }
    
    user_message = {
        "role": "user",
        "content": f"Please parse this codebook and convert it to the standardized format:\n\n{raw_text}"
    }
    
    result = generate_completion(
        client=client,
        model=model,
        messages=[system_prompt, user_message]
    )
    
    if result:
        try:
            # First try to extract JSON from markdown code blocks
            code_block_match = re.search(r'```(?:json)?\s*\n([\s\S]*?)\n```', result)
            if code_block_match:
                json_text = code_block_match.group(1)
            else:
                # If no code block, try to extract raw JSON
                json_match = re.search(r'\{[\s\S]*\}', result)
                if json_match:
                    json_text = json_match.group(0)
                else:
                    st.error("Could not extract JSON from AI response")
                    return None
            
            # Remove any trailing commas before closing brackets/braces
            json_text = re.sub(r',\s*(\]|\})', r'\1', json_text)
            # Remove // comments
            json_text = re.sub(r'//.*$', '', json_text, flags=re.MULTILINE)
            
            return json.loads(json_text)
        except json.JSONDecodeError as e:
            st.error(f"Error parsing JSON response: {e}")
            st.text("Extracted JSON text:")
            st.code(json_text, language='json')
            st.text("Full AI Response:")
            st.text(result)
            return None
    return None

# ==================== CONVERSION UTILITIES ====================

def convert_to_markdown(codebook_dict: Dict) -> str:
    """Convert the standardized codebook dictionary to markdown format."""
    lines = []
    
    # Add title if present
    if codebook_dict.get('title'):
        lines.append(f"# {codebook_dict['title']}")
        lines.append("")
    
    # Process parent codes
    for parent in codebook_dict.get('parent_codes', []):
        lines.append(f"## [{parent['id']}] {parent['name']}")
        lines.append("")
        
        # Add parent metadata
        if parent.get('definition'):
            lines.append(f"**Definition:** {parent['definition']}")
            lines.append("")
        if parent.get('inclusion_criteria'):
            lines.append(f"**Inclusion criteria:** {parent['inclusion_criteria']}")
            lines.append("")
        if parent.get('exclusion_criteria'):
            lines.append(f"**Exclusion criteria:** {parent['exclusion_criteria']}")
            lines.append("")
        if parent.get('examples'):
            lines.append(f"**Examples:** {parent['examples']}")
            lines.append("")
        
        # Process child codes
        for child in parent.get('children', []):
            lines.append(f"### [{child['id']}] {child['name']}")
            lines.append("")
            
            # Add child metadata
            if child.get('definition'):
                lines.append(f"**Definition:** {child['definition']}")
                lines.append("")
            if child.get('inclusion_criteria'):
                lines.append(f"**Inclusion criteria:** {child['inclusion_criteria']}")
                lines.append("")
            if child.get('exclusion_criteria'):
                lines.append(f"**Exclusion criteria:** {child['exclusion_criteria']}")
                lines.append("")
            if child.get('examples'):
                lines.append(f"**Examples:** {child['examples']}")
                lines.append("")
            
            # Process grandchild codes
            for grandchild in child.get('children', []):
                lines.append(f"#### [{grandchild['id']}] {grandchild['name']}")
                lines.append("")
                
                # Add grandchild metadata
                if grandchild.get('definition'):
                    lines.append(f"**Definition:** {grandchild['definition']}")
                    lines.append("")
                if grandchild.get('inclusion_criteria'):
                    lines.append(f"**Inclusion criteria:** {grandchild['inclusion_criteria']}")
                    lines.append("")
                if grandchild.get('exclusion_criteria'):
                    lines.append(f"**Exclusion criteria:** {grandchild['exclusion_criteria']}")
                    lines.append("")
                if grandchild.get('examples'):
                    lines.append(f"**Examples:** {grandchild['examples']}")
                    lines.append("")
                
                # Process great grandchild codes
                for great_grandchild in grandchild.get('children', []):
                    lines.append(f"##### [{great_grandchild['id']}] {great_grandchild['name']}")
                    lines.append("")
                    
                    # Add great grandchild metadata
                    if great_grandchild.get('definition'):
                        lines.append(f"**Definition:** {great_grandchild['definition']}")
                        lines.append("")
                    if great_grandchild.get('inclusion_criteria'):
                        lines.append(f"**Inclusion criteria:** {great_grandchild['inclusion_criteria']}")
                        lines.append("")
                    if great_grandchild.get('exclusion_criteria'):
                        lines.append(f"**Exclusion criteria:** {great_grandchild['exclusion_criteria']}")
                        lines.append("")
                    if great_grandchild.get('examples'):
                        lines.append(f"**Examples:** {great_grandchild['examples']}")
                        lines.append("")
    
    return '\n'.join(lines)

def display_codebook_structure(codebook_dict: Dict):
    """Display the codebook structure in a hierarchical view."""
    st.subheader("Detected Structure")
    
    if codebook_dict.get('title'):
        st.markdown(f"**Title:** {codebook_dict['title']}")
    
    # Display parent codes
    for parent in codebook_dict.get('parent_codes', []):
        with st.expander(f"📁 {parent['id']}: {parent['name']}"):
            if parent.get('definition'):
                st.markdown(f"**Definition:** {parent['definition']}")
            
            # Display child codes
            for child in parent.get('children', []):
                st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;📄 **{child['id']}:** {child['name']}")
                
                # Display grandchild codes
                for grandchild in child.get('children', []):
                    st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;📄 **{grandchild['id']}:** {grandchild['name']}")
                    
                    # Display great grandchild codes
                    for great_grandchild in grandchild.get('children', []):
                        st.markdown(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;📄 **{great_grandchild['id']}:** {great_grandchild['name']}")

# ==================== MAIN APPLICATION ====================

def main():
    """Main application entry point."""
    st.set_page_config(
        page_title="Codebook Generalizer",
        page_icon="📚",
        layout="wide"
    )
    
    st.title("📚 Codebook Generalizer")
    st.markdown("""
    This tool converts various qualitative codebook formats into a standardized hierarchical structure.
    
    **Supported formats:**
    - Theme → Code structures
    - Parent/Child hierarchies
    - Numbered codes (1.1, 1.1.1)
    - Table-based codebooks
    """)
    
    # Initialize session state
    if 'api_params' not in st.session_state:
        st.session_state.api_params = load_api_params()
    
    if 'client' not in st.session_state:
        st.session_state.client = initialize_client(st.session_state.api_params)
    
    # File upload
    uploaded_file = st.file_uploader(
        "Upload your codebook",
        type=['txt', 'docx', 'pdf'],
        help="Upload a codebook in any format"
    )
    
    if uploaded_file is not None:
        col1, col2 = st.columns([1, 1])
        
        with col1:
            st.subheader("Original Content")
            
            # Parse the file
            raw_text = parse_file(uploaded_file)
            
            if raw_text:
                # Clean the text
                cleaned_text = clean_text(raw_text)
                
                # Display in a scrollable container
                st.text_area(
                    "Raw codebook content:",
                    value=cleaned_text,
                    height=400,
                    disabled=True
                )
                
                # Process button
                if st.button("🔄 Process Codebook", type="primary"):
                    with st.spinner("Analyzing codebook structure..."):
                        # Parse with AI
                        parsed_codebook = detect_and_parse_codebook(
                            st.session_state.client,
                            st.session_state.api_params['MODEL'],
                            cleaned_text
                        )
                        
                        if parsed_codebook:
                            st.session_state.parsed_codebook = parsed_codebook
                            st.session_state.markdown_output = convert_to_markdown(parsed_codebook)
                            st.success("Codebook processed successfully!")
                        else:
                            st.error("Failed to parse codebook")
        
        with col2:
            if 'parsed_codebook' in st.session_state:
                st.subheader("Standardized Output")
                
                # Display structure
                display_codebook_structure(st.session_state.parsed_codebook)
                
                # Show markdown output
                with st.expander("View Markdown Output"):
                    st.text_area(
                        "Standardized markdown:",
                        value=st.session_state.markdown_output,
                        height=400,
                        disabled=True
                    )
                
                # Download options
                st.subheader("Download Options")
                
                col1_dl, col2_dl = st.columns(2)
                
                with col1_dl:
                    # Download as markdown
                    st.download_button(
                        label="📄 Download as Markdown",
                        data=st.session_state.markdown_output,
                        file_name="standardized_codebook.txt",
                        mime="text/plain"
                    )
                
                with col2_dl:
                    # Download as JSON
                    st.download_button(
                        label="📊 Download as JSON",
                        data=json.dumps(st.session_state.parsed_codebook, indent=2),
                        file_name="standardized_codebook.json",
                        mime="application/json"
                    )
                
                # Show JSON structure
                with st.expander("View JSON Structure"):
                    st.json(st.session_state.parsed_codebook)

if __name__ == "__main__":
    main()
